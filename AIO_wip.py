'''
status: WIP 
1) to test code
2) to somehow integrate HAZOP knowledge into Knowledge graph or as LLM query questions
pip install pywin32 neo4j pandas docx
'''
import win32com.client
import pandas as pd

# Connect to HYSYS
hysys = win32com.client.Dispatch("HYSYS.Application")
case = hysys.ActiveDocument

# Extract unit operations data
unit_ops = []
for unit in case.Flowsheet.Operations:
    unit_ops.append({
        "Name": unit.Name,
        "Type": unit.Type,
        "Description": unit.Description
    })

# Extract streams data
streams = []
for stream in case.Flowsheet.MaterialStreams:
    streams.append({
        "Name": stream.Name,
        "Temperature": stream.TemperatureValue,
        "Pressure": stream.PressureValue,
        "Flow": stream.MassFlowValue
    })

# Save data to CSV
unit_ops_df = pd.DataFrame(unit_ops)
streams_df = pd.DataFrame(streams)
unit_ops_df.to_csv('unit_operations.csv', index=False)
streams_df.to_csv('streams.csv', index=False)



from neo4j import GraphDatabase

# Neo4j connection details
uri = "bolt://localhost:7687"
user = "neo4j"
password = "password"

# Connect to Neo4j
driver = GraphDatabase.driver(uri, auth=(user, password))

def create_unit_op(tx, name, type, description):
    tx.run("MERGE (u:UnitOperation {name: $name, type: $type, description: $description})",
           name=name, type=type, description=description)

def create_stream(tx, name, temperature, pressure, flow):
    tx.run("MERGE (s:Stream {name: $name, temperature: $temperature, pressure: $pressure, flow: $flow})",
           name=name, temperature=temperature, pressure=pressure, flow=flow)

# Read CSV and create graph
with driver.session() as session:
    unit_ops_df = pd.read_csv('unit_operations.csv')
    streams_df = pd.read_csv('streams.csv')
    
    for _, row in unit_ops_df.iterrows():
        session.write_transaction(create_unit_op, row['Name'], row['Type'], row['Description'])
    
    for _, row in streams_df.iterrows():
        session.write_transaction(create_stream, row['Name'], row['Temperature'], row['Pressure'], row['Flow'])

driver.close()



import openai
from docx import Document

# Set up OpenAI API
openai.api_key = 'your_openai_api_key'

# Query Neo4j to get data for HAZOP
def query_neo4j(query):
    with driver.session() as session:
        result = session.run(query)
        return [record for record in result]

# Generate HAZOP dataset
hazop_data = query_neo4j("MATCH (u:UnitOperation)-[r]->(s:Stream) RETURN u, r, s")

# Generate HAZOP report using LLaMA 3.2
response = openai.Completion.create(
    model="text-davinci-003",
    prompt=f"Generate a HAZOP report based on the following data: {hazop_data}",
    max_tokens=1500
)

# Save HAZOP dataset to CSV
hazop_df = pd.DataFrame(hazop_data)
hazop_df.to_csv('hazop_dataset.csv', index=False)

# Save HAZOP report to DOCX
doc = Document()
doc.add_heading('HAZOP Report', 0)
doc.add_paragraph(response.choices[0].text)
doc.save('hazop_report.docx')
